import io

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.section import Section
from docx.shared import Mm, Pt
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

from .utils.optimize_lyrics import get_optimized_font_size


def make_picmuvie_doc(song) -> bytes:
    doc = make_doc(song.title, song.lyrics)
    return convert_docs_to_bytes(doc)


def make_doc(title: str, lyrics: list[str]) -> Document:
    doc = DocxTemplate("./picture_muvie/templates/title.docx")
    make_title_page(doc, title)
    doc.add_section()
    set_a4_landscape_section(doc, 1)
    font_size = get_optimized_font_size(lyrics)

    for i in range(len(lyrics)):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_text(paragraph, lyrics[i], "malgun_gothic", font_size, True)
        paragraph_page_number = doc.add_paragraph()
        paragraph_page_number.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_formatted_text(
            paragraph_page_number, f"{i + 1}", "malgun_gothic", 10, False
        )
    doc.add_paragraph()

    return doc


def convert_docs_to_bytes(doc: Document):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def set_section_margin(sec: Section, top: int, bottom: int, left: int, right: int):
    # margin setting
    sec.top_margin = Mm(top)
    sec.bottom_margin = Mm(bottom)
    sec.left_margin = Mm(left)
    sec.right_margin = Mm(right)


def set_a4_landscape_section(doc: Document, sect_index: int):
    sections = doc.sections
    section = sections[sect_index]
    # a4 size in 'mm' unit
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # portrait -> landscape
    section.page_width, section.page_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    set_section_margin(section, 165, 20, 24, 24)


"""
HWP 파일에만 적용되지 않아 추후 보완 후 적용 위해 주석 처리

def set_page_number_in_section(doc: Document, sect_index: int):
    sections = doc.sections
    sec = sections[sect_index]
    header = sec.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(f"{sect_index}")
    run.font.size = Pt(15)
"""


def add_formatted_text(
    paragraph: Paragraph, text: str, font_name: str, size: int, is_bold: bool
):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = is_bold


def make_title_page(doc: Document, title: str):
    context = {
        "title": title,
        "school": "OO초등학교",
        "class": "O학년 O반",
    }
    doc.render(context)
